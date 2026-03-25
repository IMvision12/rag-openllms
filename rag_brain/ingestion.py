from pathlib import Path

from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from rag_brain.config import ChunkingStrategy, Settings

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def load_pdf_as_documents(path: Path | str, *, source_name: str | None = None) -> list[Document]:
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(p)
    loader = PyPDFLoader(str(p))
    docs = loader.load()
    src = source_name or p.name
    for d in docs:
        d.metadata.setdefault("source", src)
    return docs


def load_docx_as_documents(path: Path | str, *, source_name: str | None = None) -> list[Document]:
    from docx import Document as DocxDocument

    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(p)
    doc = DocxDocument(str(p))
    src = source_name or p.name
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    full_text = "\n\n".join(paragraphs)
    return [Document(page_content=full_text, metadata={"source": src})]


def load_documents(path: Path | str, *, source_name: str | None = None) -> list[Document]:
    """Dispatch loader based on file extension. Supports PDF and DOCX."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".pdf":
        return load_pdf_as_documents(p, source_name=source_name)
    if ext == ".docx":
        return load_docx_as_documents(p, source_name=source_name)
    raise ValueError(f"Unsupported file type '{ext}'. Supported: {SUPPORTED_EXTENSIONS}")


def _fixed_split(documents: list[Document], settings: Settings) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        add_start_index=True,
    )
    return splitter.split_documents(documents)


def _semantic_split(documents: list[Document], settings: Settings) -> list[Document]:
    """Sentence-boundary aware splitting: split on sentences first, then merge up to chunk_size."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", "? ", "! ", "; ", ", ", " ", ""],
        add_start_index=True,
    )
    return splitter.split_documents(documents)


def split_documents(documents: list[Document], settings: Settings) -> list[Document]:
    strategy = settings.chunking_strategy
    if strategy == ChunkingStrategy.semantic:
        return _semantic_split(documents, settings)
    return _fixed_split(documents, settings)
